import praw
import time
from docx import Document

# --- Reddit API Credentials ---
reddit = praw.Reddit(
    client_id='',
    client_secret='',
    user_agent=''
)

# --- Keywords to search for ---
keywords = [
    "Unreal Engine", "Unreal Engine 4", "Unreal Engine 5", "Epic Games", "Unreal Engine graphics", "real-time rendering", "photorealistic graphics", "Unreal Engine VR", "Unreal Engine AR", "immersive experiences", "Unreal Editor", "Unreal Marketplace", "Unreal Engine Blueprints", "Unreal Engine C++", "Unreal Engine revenue", "Unreal Engine licensing", "Unreal Engine royalties", "Unreal Engine security", "Unreal Engine optimization", "Unreal Engine performance", "Unreal Engine vs Unity", "Unreal Engine vs CryEngine", "Unreal Engine vs Godot", "Unreal Engine vs Lumberyard"
]

TARGET_URL_COUNT = 150
url_set = set()

# Exclude URLs from these external domains (e.g., video sites)
exclusion_domains = ["youtube.com", "vimeo.com"]

print("Searching across r/all using multiple keywords...")

# Loop through each keyword and search in r/all
for keyword in keywords:
    print(f"\n🔍 Searching for keyword: '{keyword}'")
    try:
        for submission in reddit.subreddit("all").search(keyword, limit=None, time_filter="all"):
            # Inclusion: Title must contain the keyword (case-insensitive)
            if keyword.upper() not in submission.title.upper():
                continue
            # Inclusion: Minimum score threshold
            if submission.score < 1:
                continue
            # Inclusion: Must not be marked NSFW
            if submission.over_18:
                continue
            # Exclusion: Skip posts with generic question titles (ending with a '?')
            if submission.title.strip().endswith("?"):
                continue
            # Exclusion: Skip posts whose URLs include excluded domains
            if any(domain in submission.url for domain in exclusion_domains):
                continue
            # NEW: Only include URLs that are from Reddit
            if "reddit.com" not in submission.url.lower():
                continue
            # Add URL if not already collected
            if submission.url not in url_set:
                url_set.add(submission.url)
                print(submission.url)
            # Break out if we reach our target number of URLs
            if len(url_set) >= TARGET_URL_COUNT:
                break
        if len(url_set) >= TARGET_URL_COUNT:
            break
        time.sleep(1)  # Delay to avoid rate limiting
    except Exception as e:
        print(f"❌ Error searching for '{keyword}': {e}")
    if len(url_set) >= TARGET_URL_COUNT:
        break

# Save collected URLs to a Word document
doc = Document()
doc.add_heading('QNX Reddit URLs', level=1)

for url in url_set:
    doc.add_paragraph(url)

output_file = "unreal_urls.docx"
doc.save(output_file)

print(f"\n✅ Successfully collected {len(url_set)} URLs and saved to '{output_file}'")
