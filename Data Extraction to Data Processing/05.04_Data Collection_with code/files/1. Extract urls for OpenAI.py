import praw, time, json, os
from prawcore.exceptions import RequestException
from docx import Document

# --- Reddit API Credentials (Set as Environment Variables for Security) ---
reddit = praw.Reddit(
    client_id='',
    client_secret='',
    user_agent=''
)

keywords = [
    "OpenAI", "OpenAI API", "OpenAI Playground", "OpenAI Developer",
    "OpenAI API Key", "OpenAI fine-tuning", "OpenAI API usage", "OpenAI API Python",
    "OpenAI GPT Models", "OpenAI GPT-3", "OpenAI GPT-4", "OpenAI Dev Tools", 
    "OpenAI Custom Models", "OpenAI Fine-Tuning Process", "OpenAI Deployment",
    "OpenAI Programming", "OpenAI Integration", "OpenAI CLI", "OpenAI SDK", 
    "OpenAI API Documentation", "OpenAI API Errors", "OpenAI Rate Limits",
    "OpenAI Prompt Engineering", "OpenAI Model Training", "OpenAI Security",
    "OpenAI Authentication", "OpenAI Scalability", "OpenAI Pricing", "OpenAI API Limits",
    "OpenAI Usage Analysis", "OpenAI Data Processing", "OpenAI Model Deployment"
]

TARGET_URL_COUNT = 400
url_set = set()

# Exclude any subreddit or URL that matches these patterns
exclusion_subreddits = {"OpenAI"}
exclusion_domains = ["youtube.com", "vimeo.com"]
exclusion_keywords = ["Ghibli", "Fortnite"]

# Required subreddit to always include
required_subreddits = {"QNX"}

def is_valid_submission(submission):
    """
    Validates if the submission should be collected based on defined filters.
    """
    if submission.over_18 or submission.score < 1:
        return False
    if any(domain in submission.url for domain in exclusion_domains):
        return False
    if any(keyword in submission.url for keyword in exclusion_keywords):
        return False
    if submission.subreddit.display_name in exclusion_subreddits:
        return False
    return True

print("Searching across r/all using multiple keywords...")

# --- Step 1: Always collect from 'QNX' subreddit ---
try:
    for submission in reddit.subreddit("QNX").new(limit=None):  # Fetching all new posts in QNX subreddit
        if is_valid_submission(submission):
            url_set.add(submission.url)
            print(f"Collected from QNX: {submission.url}")
        if len(url_set) >= TARGET_URL_COUNT:
            break
    time.sleep(1)
except RequestException:
    print("Rate limit reached, retrying after 60 seconds...")
    time.sleep(60)
except Exception as e:
    print(f"❌ Error searching in QNX subreddit: {e}")

# --- Step 2: Search across r/all with keywords ---
for keyword in keywords:
    try:
        for submission in reddit.subreddit("all").search(keyword, limit=None, time_filter="all"):
            if (keyword.upper() in submission.title.upper() and is_valid_submission(submission) 
                and not submission.title.strip().endswith("?") and "reddit.com" in submission.url.lower()):
                
                url_set.add(submission.url)
                print(f"Collected from r/all: {submission.url}")
            
            if len(url_set) >= TARGET_URL_COUNT:
                break
        time.sleep(1)
    except RequestException:
        print("Rate limit reached, retrying after 60 seconds...")
        time.sleep(60)
    except Exception as e:
        print(f"❌ Error searching for '{keyword}': {e}")
    if len(url_set) >= TARGET_URL_COUNT:
        break

# --- Save URLs to a Word Document ---
doc = Document()
doc.add_heading('QNX Reddit URLs', level=1)
for url in url_set: 
    doc.add_paragraph(url)
doc.save("qnx_urls2.docx")

# --- Save URLs to a JSON file ---
with open("qnx_urls2.json", "w") as f:
    json.dump(list(url_set), f)

print(f"\n✅ Successfully collected {len(url_set)} URLs and saved to 'qnx_urls2.docx' and 'qnx_urls2.json'")
